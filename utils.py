# utils.py
import streamlit as st
import hashlib
import os
import fitz  # PyMuPDF
import base64
import json
import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime

# -------------------------
# Session management
# -------------------------
def init_session():
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'username' not in st.session_state:
        st.session_state.username = None
    if 'user_type' not in st.session_state:
        st.session_state.user_type = None
    if 'uploaded_resumes' not in st.session_state:
        st.session_state.uploaded_resumes = {}  # {username: [file_names]}
    if 'scores' not in st.session_state:
        st.session_state.scores = {}  # {filename: (score, skills)}
    if 'dataset_files' not in st.session_state:
        st.session_state.dataset_files = {}  # {company_name: file_path}
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'theme' not in st.session_state:
        st.session_state.theme = 'light'

# -------------------------
# Password hashing
# -------------------------
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def check_password(stored_hash, password):
    return stored_hash == hash_password(password)

# -------------------------
# Theme CSS
# -------------------------
def theme_css():
    if st.session_state.theme == "dark":
        st.markdown("""
        <style>
        body {background-color:#222;color:#eee;font-family: 'Arial', sans-serif;}
        .stButton button {background-color:#444;color:#fff;}
        .stFileUploader>div>div>input {color:#fff;}
        </style>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <style>
        body {background-color:#fff;color:#000;font-family: 'Arial', sans-serif;}
        .stButton button {background-color:#eee;color:#000;}
        .stFileUploader>div>div>input {color:#000;}
        </style>
        """, unsafe_allow_html=True)

# -------------------------
# Load logo
# -------------------------
def load_logo():
    logo_path = "assets/logo.png"
    if os.path.exists(logo_path):
        return logo_path
    return None

# -------------------------
# Extract text from PDF/DOCX
# -------------------------
def extract_text(file):
    try:
        if file.type == "application/pdf":
            doc = fitz.open(stream=file.read(), filetype="pdf")
            text = "".join([page.get_text() for page in doc])
            return text
        else:
            import docx
            doc = docx.Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])
            return text
    except Exception as e:
        st.error(f"Error extracting text: {e}")
        return ""

# -------------------------
# Extract skills from text
# -------------------------
def extract_skills_from_text(text, skill_set=None):
    if skill_set is None:
        skill_set = ["Python", "Java", "SQL", "Excel", "Power BI", "Machine Learning"]
    skills = [skill for skill in skill_set if skill.lower() in text.lower()]
    return skills

# -------------------------
# Score resume (ATS)
# -------------------------
def score_resume(resume_text, required_skills):
    skills_found = extract_skills_from_text(resume_text, required_skills)
    if required_skills:
        score = len(skills_found) / len(required_skills) * 100
    else:
        score = 0
    return round(score), skills_found

# -------------------------
# Show ATS pie chart
# -------------------------
def show_score_pie(score):
    fig, ax = plt.subplots()
    ax.pie([score, 100-score], labels=["Matched", "Remaining"], colors=["#4CAF50","#ddd"], autopct='%1.1f%%')
    st.pyplot(fig)

# -------------------------
# Save resume per user
# -------------------------
def save_resume(file, username):
    folder = os.path.join("uploaded_resumes", username)
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, file.name)
    with open(path, "wb") as f:
        f.write(file.getbuffer())
    if username not in st.session_state.uploaded_resumes:
        st.session_state.uploaded_resumes[username] = []
    if file.name not in st.session_state.uploaded_resumes[username]:
        st.session_state.uploaded_resumes[username].append(file.name)
    return path

# -------------------------
# List resumes for a user
# -------------------------
def list_uploaded_resumes(username):
    return st.session_state.uploaded_resumes.get(username, [])

# -------------------------
# Preview resume (PDF/DOCX)
# -------------------------
def preview_resume(username, filename):
    folder = os.path.join("uploaded_resumes", username)
    path = os.path.join(folder, filename)
    if filename.lower().endswith(".pdf"):
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode('utf-8')
    else:
        import docx
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])

# -------------------------
# Chat message send/receive
# -------------------------
def send_message(sender, message):
    st.session_state.chat_history.append({
        "sender": sender,
        "message": message,
        "time": datetime.now().strftime("%H:%M:%S")
    })

# -------------------------
# Save dataset CSV for company
# -------------------------
def save_dataset(file, company_name):
    folder = "uploaded_datasets"
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, file.name)
    with open(path, "wb") as f:
        f.write(file.getbuffer())
    st.session_state.dataset_files[company_name] = path
    return path

# -------------------------
# Search resumes by keyword
# -------------------------
def search_resumes(keyword):
    results = []
    for username, files in st.session_state.uploaded_resumes.items():
        for resume_name in files:
            path = os.path.join("uploaded_resumes", username, resume_name)
            text = extract_text(open(path, "rb"))
            if keyword.lower() in text.lower():
                results.append((username, resume_name))
    return results

# -------------------------
# Rank resumes by ATS score
# -------------------------
def rank_resumes():
    ranked = sorted(st.session_state.scores.items(), key=lambda x: x[1][0], reverse=True)
    return ranked
